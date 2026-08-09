"""Native (non-OCR) text extraction from PDF/DOCX, and input-file discovery."""

import logging
import os
import sys
from dataclasses import dataclass, field
from pathlib import Path

from hollyocr.utils.logging_setup import LOGGER

try:
    from pypdf import PdfReader
    from pypdf.generic import ContentStream
except Exception:
    PdfReader = None
    ContentStream = None

try:
    import docx
except Exception:
    docx = None

ENABLE_PYMUPDF_ENV = "HOLLYOCR_ENABLE_PYMUPDF"
MUPDF_DISABLED_FOR_STABILITY = (
    os.environ.get("HOLLYOCR_DISABLE_PYMUPDF", "").strip().lower() in {"1", "true", "yes", "on"}
) or (
    sys.platform == "darwin"
    and getattr(sys, "frozen", False)
    and (os.environ.get(ENABLE_PYMUPDF_ENV, "").strip().lower() not in {"1", "true", "yes", "on"})
)
_mupdf_disable_notice_emitted = False


@dataclass(frozen=True)
class PdfInspection:
    """Cheap structural facts used to choose the safest extraction path."""

    page_count: int
    image_pages: frozenset
    image_counts: dict
    reliable: bool = True
    warning: str = ""
    image_coverages: dict = field(default_factory=dict)
    coverage_reliable: bool = True

if MUPDF_DISABLED_FOR_STABILITY:
    fitz = None
    pymupdf4llm = None
else:
    try:
        import pymupdf as fitz
    except Exception:
        fitz = None

    try:
        import pymupdf4llm
    except Exception:
        pymupdf4llm = None


def find_files(input_path):
    """Accept either a single file path, list of file paths, or directory path."""
    valid_extensions = {'.pdf', '.md', '.markdown', '.docx', '.png', '.jpg', '.jpeg'}
    # If input_path is a list (multiple files selected)
    if isinstance(input_path, list):
        for f in input_path:
            p = Path(f)
            if p.is_file() and p.suffix.lower() in valid_extensions:
                yield p
    else:
        # Single path (file or directory)
        p = Path(input_path)
        if p.is_file() and p.suffix.lower() in valid_extensions:
            yield p
        elif p.is_dir():
            files = (
                f for f in p.rglob('*')
                if f.is_file() and f.suffix.lower() in valid_extensions
            )
            for f in sorted(files, key=lambda item: str(item).lower()):
                yield f
        else:
            return


def _normalize_block_text(text):
    lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
    return "\n".join(lines).strip()


def _sort_pdf_blocks_for_reading(blocks, page_width):
    """Approximate reading order for single- and double-column documents."""
    if not blocks:
        return []

    sorted_by_x = sorted(blocks, key=lambda block: block["x0"])
    largest_gap = 0.0
    gap_index = None
    for idx in range(1, len(sorted_by_x)):
        gap = sorted_by_x[idx]["x0"] - sorted_by_x[idx - 1]["x0"]
        if gap > largest_gap:
            largest_gap = gap
            gap_index = idx

    if (
        gap_index is not None
        and page_width > 0
        and largest_gap >= (page_width * 0.18)
        and gap_index >= 2
        and (len(sorted_by_x) - gap_index) >= 2
    ):
        left_blocks = sorted(sorted_by_x[:gap_index], key=lambda block: (round(block["y0"], 1), round(block["x0"], 1)))
        right_blocks = sorted(sorted_by_x[gap_index:], key=lambda block: (round(block["y0"], 1), round(block["x0"], 1)))
        return left_blocks + right_blocks

    return sorted(blocks, key=lambda block: (round(block["y0"], 1), round(block["x0"], 1)))


def extract_fitz_page_text(page):
    """Extract text using positioned text blocks to preserve reading order better than plain get_text()."""
    try:
        raw_blocks = page.get_text("blocks")
    except Exception:
        return page.get_text() or ""

    blocks = []
    page_width = float(getattr(page.rect, "width", 0.0) or 0.0)
    for block in raw_blocks or []:
        if len(block) < 5:
            continue
        x0, y0, x1, y1, text = block[:5]
        block_type = block[6] if len(block) > 6 else 0
        if block_type not in (0, None):
            continue
        cleaned = _normalize_block_text(text)
        if not cleaned:
            continue
        blocks.append({
            "x0": float(x0),
            "y0": float(y0),
            "x1": float(x1),
            "y1": float(y1),
            "text": cleaned,
        })

    ordered = _sort_pdf_blocks_for_reading(blocks, page_width)
    return "\n\n".join(block["text"] for block in ordered).strip()


class _PdfPageWarningCollector(logging.Handler):
    def __init__(self):
        super().__init__(level=logging.WARNING)
        self.messages = []

    def emit(self, record):
        self.messages.append(record.getMessage())


def _extract_pypdf_page_text(page, page_index, diagnostic_callback=None):
    """Extract one pypdf page and surface warnings that imply lost Form text."""
    collector = _PdfPageWarningCollector()
    page_logger = logging.getLogger("pypdf._page")
    page_logger.addHandler(collector)
    try:
        text = page.extract_text() or ""
    except Exception as exc:
        if diagnostic_callback:
            diagnostic_callback(page_index, f"Falha na extração nativa: {exc}")
        return ""
    finally:
        page_logger.removeHandler(collector)

    form_failures = [
        message.strip()
        for message in collector.messages
        if "impossible to decode XFormObject" in message
    ]
    if form_failures and diagnostic_callback:
        diagnostic_callback(
            page_index,
            "Camada nativa incompleta: " + "; ".join(dict.fromkeys(form_failures)),
        )
    return text


def iter_extracted_text_pages(
    pdf_path: Path,
    output_format='txt',
    use_ocr=False,
    lang='por',
    allow_markdown_layout=True,
    diagnostic_callback=None,
):
    """Yield text per page using the best available native extractor."""
    output_format = (output_format or 'txt').lower().strip()
    global _mupdf_disable_notice_emitted
    if MUPDF_DISABLED_FOR_STABILITY and not _mupdf_disable_notice_emitted:
        LOGGER.warning(
            "PyMuPDF desativado neste app macOS empacotado para estabilidade. "
            "Use %s=1 para reativar manualmente.",
            ENABLE_PYMUPDF_ENV,
        )
        _mupdf_disable_notice_emitted = True

    # Try pymupdf4llm for markdown format
    if output_format == 'md' and allow_markdown_layout and pymupdf4llm is not None:
        try:
            try:
                chunks = pymupdf4llm.to_markdown(
                    str(pdf_path),
                    page_chunks=True,
                    # OCR is handled by this pipeline so Tesseract path, language and progress stay consistent.
                    use_ocr=False,
                    ocr_language=lang or 'por',
                )
            except TypeError:
                chunks = pymupdf4llm.to_markdown(str(pdf_path), page_chunks=True)
            for chunk in chunks:
                yield chunk.get("text", "")
            return
        except Exception as e:
            LOGGER.exception("pymupdf4llm failed for %s", pdf_path)
            print(f"pymupdf4llm failed, falling back to pymupdf: {e}")

    # Try PyMuPDF (fitz)
    if fitz is not None:
        try:
            with fitz.open(str(pdf_path)) as doc:
                for page_index, page in enumerate(doc):
                    try:
                        txt = extract_fitz_page_text(page)
                    except Exception as exc:
                        txt = ''
                        if diagnostic_callback:
                            diagnostic_callback(page_index, f"Falha na extração nativa: {exc}")
                    yield txt
            return
        except Exception as e:
            LOGGER.exception("PyMuPDF failed to extract text from %s", pdf_path)
            print(f"PyMuPDF failed to extract text: {e}. Falling back to pypdf.")

    # Fallback to pypdf
    try:
        if PdfReader is None:
            raise RuntimeError("pypdf não está disponível para fallback de texto nativo.")
        with pdf_path.open("rb") as pdf_stream:
            reader = PdfReader(pdf_stream)
            for page_index, page in enumerate(reader.pages):
                yield _extract_pypdf_page_text(page, page_index, diagnostic_callback)
        return
    except Exception as e:
        LOGGER.exception("pypdf failed to extract text from %s", pdf_path)
        print(f"pypdf failed to extract text: {e}")
        return


def _resolve_pdf_object(value):
    getter = getattr(value, "get_object", None)
    return getter() if callable(getter) else value


def _count_images_in_resources(resources, visited=None):
    """Count image XObjects, including images nested inside Form XObjects."""
    resources = _resolve_pdf_object(resources)
    if not resources:
        return 0
    visited = visited if visited is not None else set()
    xobjects = _resolve_pdf_object(resources.get("/XObject")) if hasattr(resources, "get") else None
    if not xobjects:
        return 0

    count = 0
    for ref in xobjects.values():
        marker = getattr(ref, "idnum", None)
        if marker is not None and marker in visited:
            continue
        if marker is not None:
            visited.add(marker)
        obj = _resolve_pdf_object(ref)
        if not hasattr(obj, "get"):
            continue
        subtype = obj.get("/Subtype")
        if subtype == "/Image":
            count += 1
        elif subtype == "/Form":
            count += _count_images_in_resources(obj.get("/Resources"), visited)
    return count


def _matrix_area_scale(matrix):
    """Return the absolute area scale of a six-value PDF transform matrix."""
    try:
        a, b, c, d = (float(value) for value in list(matrix)[:4])
        return abs((a * d) - (b * c))
    except Exception:
        return 1.0


def _image_coverages_in_stream(
    stream,
    resources,
    reader,
    page_area,
    initial_scale=1.0,
    active_forms=None,
    form_cache=None,
    depth=0,
):
    """Measure every drawn image as a fraction of the visible page area.

    PDF image XObjects occupy a unit square transformed by the current graphics
    matrix.  The determinant of that matrix therefore gives the drawn area in
    page units.  Form XObjects are followed recursively because court PDFs
    commonly wrap logos, signatures and scanned pages inside forms.
    """
    if ContentStream is None or not stream or depth > 12 or page_area <= 0:
        return []
    resources = _resolve_pdf_object(resources)
    if not resources:
        return []

    content = ContentStream(_resolve_pdf_object(stream), reader)
    scale = float(initial_scale or 0.0)
    stack = []
    coverages = []
    active_forms = active_forms if active_forms is not None else set()
    form_cache = form_cache if form_cache is not None else {}

    for operands, operator in content.operations:
        if operator == b"q":
            stack.append(scale)
            continue
        if operator == b"Q":
            if stack:
                scale = stack.pop()
            continue
        if operator == b"cm" and len(operands) >= 6:
            scale *= _matrix_area_scale(operands)
            continue
        if operator != b"Do" or not operands:
            continue

        xobjects = _resolve_pdf_object(resources.get("/XObject")) if hasattr(resources, "get") else None
        if not xobjects:
            continue
        ref = xobjects.get(operands[0]) if hasattr(xobjects, "get") else None
        if ref is None:
            continue
        obj = _resolve_pdf_object(ref)
        if not hasattr(obj, "get"):
            continue
        subtype = obj.get("/Subtype")
        if subtype == "/Image":
            # Clamp malformed/off-page transforms so one bad object does not
            # distort the automatic OCR strategy.
            coverages.append(max(0.0, min(1.0, scale / page_area)))
            continue
        if subtype != "/Form":
            continue

        marker = getattr(ref, "idnum", None)
        marker = ("ref", marker) if marker is not None else ("obj", id(obj))
        if marker in active_forms:
            continue
        active_forms.add(marker)
        try:
            form_matrix = obj.get("/Matrix") or [1, 0, 0, 1, 0, 0]
            form_resources = obj.get("/Resources") or resources
            cache_key = (marker, round(float(page_area), 3))
            relative_coverages = form_cache.get(cache_key)
            if relative_coverages is None:
                relative_coverages = tuple(_image_coverages_in_stream(
                    obj,
                    form_resources,
                    reader,
                    page_area,
                    initial_scale=_matrix_area_scale(form_matrix),
                    active_forms=active_forms,
                    form_cache=form_cache,
                    depth=depth + 1,
                ))
                form_cache[cache_key] = relative_coverages
            coverages.extend(max(0.0, min(1.0, value * scale)) for value in relative_coverages)
        finally:
            active_forms.discard(marker)
    return coverages


def _page_image_coverages(page, reader, form_cache=None):
    resources = page.get("/Resources") if hasattr(page, "get") else None
    # Use the raw /Contents object. pypdf 6's PageObject.get_contents() may
    # normalize an array of streams to an empty ContentStream for some PDFs,
    # which loses the transformation commands needed to measure image area.
    contents = page.get("/Contents") if hasattr(page, "get") else None
    if contents is None and hasattr(page, "get_contents"):
        contents = page.get_contents()
    box = getattr(page, "cropbox", None) or getattr(page, "mediabox", None)
    try:
        page_area = abs(float(box.width) * float(box.height))
    except Exception:
        page_area = 0.0
    return _image_coverages_in_stream(contents, resources, reader, page_area, form_cache=form_cache)


def inspect_pdf(pdf_path: Path):
    """Inspect page count and embedded raster images without rendering the PDF.

    If structural inspection fails, ``reliable`` is false. The caller can then
    conservatively OCR every page, which is preferable to silently missing a
    screenshot used as evidence.
    """
    if PdfReader is None:
        return PdfInspection(0, frozenset(), {}, False, "pypdf indisponível para inspecionar imagens.")
    try:
        with pdf_path.open("rb") as pdf_stream:
            reader = PdfReader(pdf_stream)
            image_counts = {}
            image_coverages = {}
            coverage_warnings = []
            form_cache = {}
            for page_index, page in enumerate(reader.pages):
                resources = page.get("/Resources") if hasattr(page, "get") else None
                image_count = _count_images_in_resources(resources)
                if image_count:
                    image_counts[page_index] = image_count
                    try:
                        coverages = _page_image_coverages(page, reader, form_cache=form_cache)
                        if coverages:
                            image_coverages[page_index] = tuple(coverages)
                        else:
                            coverage_warnings.append(page_index + 1)
                    except Exception:
                        coverage_warnings.append(page_index + 1)
                        LOGGER.warning("Could not measure image placement on page %s of %s", page_index + 1, pdf_path)
            warning = ""
            if coverage_warnings:
                preview = ", ".join(str(page) for page in coverage_warnings[:8])
                suffix = "..." if len(coverage_warnings) > 8 else ""
                warning = f"Não foi possível medir a área das imagens nas páginas {preview}{suffix}."
            return PdfInspection(
                page_count=len(reader.pages),
                image_pages=frozenset(image_counts),
                image_counts=image_counts,
                warning=warning,
                image_coverages=image_coverages,
                coverage_reliable=not coverage_warnings,
            )
    except Exception as exc:
        LOGGER.exception("Failed to inspect PDF structure for %s", pdf_path)
        return PdfInspection(0, frozenset(), {}, False, f"Falha ao inspecionar imagens: {exc}")


def extract_text_pages(
    pdf_path: Path,
    output_format='txt',
    use_ocr=False,
    lang='por',
    allow_markdown_layout=True,
    diagnostic_callback=None,
):
    """Return a lazy iterator so very large PDFs do not occupy RAM twice."""
    return iter_extracted_text_pages(
        pdf_path,
        output_format=output_format,
        use_ocr=use_ocr,
        lang=lang,
        allow_markdown_layout=allow_markdown_layout,
        diagnostic_callback=diagnostic_callback,
    )


def extract_table_text(table):
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            value = ' '.join(cell.text.split())
            cells.append(value)
        row_text = '\t'.join(cells).strip()
        if row_text:
            rows.append(row_text)
    return '\n'.join(rows)


def extract_docx_text(file_path: Path):
    document = docx.Document(file_path)
    parts = []

    def add_text(text):
        text = text.strip()
        if text:
            parts.append(text)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            add_text(paragraph.text)

    for paragraph in document.paragraphs:
        add_text(paragraph.text)

    for table in document.tables:
        add_text(extract_table_text(table))

    for section in document.sections:
        for paragraph in section.footer.paragraphs:
            add_text(paragraph.text)

    return '\n'.join(parts)


def normalize_output_format(output_format):
    output_format = (output_format or 'txt').lower().strip()
    return output_format if output_format in {'txt', 'md'} else 'txt'


def get_pages_count(pages_source):
    if hasattr(pages_source, "page_count"):
        return int(getattr(pages_source, "page_count"))
    return len(pages_source)


def get_page_text_at(pages_source, index):
    if hasattr(pages_source, "get_text"):
        return pages_source.get_text(index)
    return pages_source[index]
