import logging
import tempfile
import unittest
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from unittest import mock

import pymupdf as fitz
from docx import Document
from PIL import Image, ImageDraw

from hollyocr import app
from hollyocr.core import extraction as core_extraction
from hollyocr.core import ocr as core_ocr
from hollyocr.core import pipeline as core_pipeline
from hollyocr.core import vision_ocr
from hollyocr.core.extraction import (
    _extract_pypdf_page_text,
    extract_text_pages,
    inspect_pdf,
)
from hollyocr.core.quality import merge_native_and_ocr_text
from hollyocr.core.strategy import choose_processing_profile
from hollyocr.gui.shared import processing as gui_shared_processing


def create_searchable_pdf(path: Path, text: str):
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 96), text, fontsize=12)
    doc.save(path)
    doc.close()


def create_mixed_pdf(path: Path):
    screenshot = path.with_suffix(".png")
    image = Image.new("RGB", (900, 420), "white")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((20, 20, 880, 400), radius=18, fill="#E7FFDB")
    draw.text((55, 70), "WhatsApp 14:32", fill="black")
    draw.text((55, 145), "Mensagem exclusiva dentro do print", fill="black")
    draw.text((55, 220), "Comprovante enviado em 20/06/2026", fill="black")
    image.save(screenshot)

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    page.insert_text((50, 55), "Texto nativo do processo que deve ser preservado.", fontsize=12)
    page.insert_image(fitz.Rect(50, 100, 545, 331), filename=str(screenshot))
    doc.save(path)
    doc.close()
    screenshot.unlink()


class PipelineMinimalTests(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory(prefix="pdf_ocr_tests_")
        self.base = Path(self.tmp.name)
        self.input_dir = self.base / "input"
        self.output_dir = self.base / "output"
        self.input_dir.mkdir()
        self.output_dir.mkdir()
        self.pdf = self.input_dir / "processo.pdf"
        self.mixed_pdf = self.input_dir / "processo_com_print.pdf"
        create_searchable_pdf(self.pdf, "Documento A pagina 1.")
        create_mixed_pdf(self.mixed_pdf)

        self.image = self.input_dir / "print.png"
        Image.new("RGB", (500, 200), "white").save(self.image)
        self.md = self.input_dir / "notas.md"
        self.md.write_text("# Titulo\n\nConteudo integral.", encoding="utf-8")
        self.docx = self.input_dir / "processo.docx"
        doc = Document()
        doc.add_heading("Titulo DOCX", 1)
        doc.add_paragraph("Texto do documento.")
        doc.save(self.docx)

    def tearDown(self):
        self.tmp.cleanup()

    def test_native_pdf_to_markdown(self):
        meta, _ = app.process_file(self.pdf, self.output_dir, self.input_dir, use_ocr=False, output_format="md")
        content = Path(meta["output_path"]).read_text(encoding="utf-8")
        self.assertEqual(meta["status"], "ok")
        self.assertIn("## Resumo da extração", content)
        self.assertIn("## Página 1", content)
        self.assertIn("Documento A", content)

    def test_image_page_is_detected_even_with_good_native_text(self):
        inspection = inspect_pdf(self.mixed_pdf)
        self.assertTrue(inspection.reliable)
        self.assertEqual(inspection.page_count, 1)
        self.assertEqual(inspection.image_pages, frozenset({0}))
        self.assertGreater(max(inspection.image_coverages[0]), 0.20)

    def test_mixed_page_runs_ocr_and_preserves_both_layers(self):
        with mock.patch.object(core_pipeline, "convert_from_path", return_value=["render.jpg"]), \
             mock.patch.object(core_pipeline, "ocr_image_file", return_value="WhatsApp 14:32\nMensagem exclusiva dentro do print"):
            meta, _ = app.process_file(
                self.mixed_pdf,
                self.output_dir,
                self.input_dir,
                use_ocr=True,
                output_format="md",
                workers=1,
            )
        content = Path(meta["output_path"]).read_text(encoding="utf-8")
        self.assertIn("Texto nativo do processo", content)
        self.assertIn("Mensagem exclusiva dentro do print", content)
        self.assertIn("Conteúdo adicional detectado por OCR", content)
        self.assertEqual(meta["page_details"][0]["source"], "nativo+ocr")
        self.assertIn("imagem relevante", meta["page_details"][0]["ocr_reason"])

    def test_force_ocr_does_not_erase_native_text(self):
        with mock.patch.object(core_pipeline, "convert_from_path", return_value=["render.jpg"]), \
             mock.patch.object(core_pipeline, "ocr_image_file", return_value="Texto visto somente pelo OCR"):
            meta, _ = app.process_file(
                self.pdf, self.output_dir, self.input_dir, use_ocr=True, force_ocr=True, output_format="md", workers=1
            )
        content = Path(meta["output_path"]).read_text(encoding="utf-8")
        self.assertIn("Documento A pagina 1", content)
        self.assertIn("Texto visto somente pelo OCR", content)

    def test_merge_deduplicates_equivalent_lines(self):
        merged = merge_native_and_ocr_text("Linha principal do processo", "Linha principal do processo\nProva no print")
        self.assertEqual(merged.count("Linha principal do processo"), 1)
        self.assertIn("Prova no print", merged)

    def test_merge_deduplicates_native_line_wrapping_and_accents(self):
        native = "Esta é uma manifest aç ão longa do processo que foi\nquebrada pela camada nativa."
        ocr = "Esta e uma manifestacao longa do processo que foi quebrada pela camada nativa.\nMensagem nova no print"
        merged = merge_native_and_ocr_text(native, ocr)
        self.assertNotIn("Esta e uma manifestacao", merged)
        self.assertIn("Mensagem nova no print", merged)

    def test_merge_discards_fragmented_ocr_when_tokens_are_already_native(self):
        native = "data lançamentos valor saldo 86,62 509,60 8,84 0,17 823,52 116,57"
        ocr = "data\nlançamentos\nvalor\nsaldo\n86,62\n509,60\n8,84\n0,17\n823,52\n116,57"
        self.assertEqual(merge_native_and_ocr_text(native, ocr), native)

    def test_image_ocr_policy_ignores_logos_and_complete_scan_layers(self):
        rich_native = {"char_count": 1800, "quality_score": 100, "needs_ocr": False}
        sparse_native = {"char_count": 250, "quality_score": 100, "needs_ocr": False}
        self.assertFalse(core_pipeline._image_page_requires_ocr(rich_native, (0.011,), True))
        self.assertTrue(core_pipeline._image_page_requires_ocr(rich_native, (0.22,), True))
        self.assertFalse(core_pipeline._image_page_requires_ocr(rich_native, (0.78,), True))
        self.assertTrue(core_pipeline._image_page_requires_ocr(sparse_native, (0.78,), True))
        self.assertTrue(core_pipeline._image_page_requires_ocr(rich_native, (), True))

    def test_pypdf_form_decode_warning_is_reported_as_extraction_issue(self):
        class DamagedFormPage:
            def extract_text(self):
                logging.getLogger("pypdf._page").warning(" impossible to decode XFormObject /Xf1")
                return "Somente o rodapé técnico"

        issues = []
        text = _extract_pypdf_page_text(
            DamagedFormPage(), 7, lambda page, message: issues.append((page, message))
        )
        self.assertEqual(text, "Somente o rodapé técnico")
        self.assertEqual(issues[0][0], 7)
        self.assertIn("Camada nativa incompleta", issues[0][1])

    def test_native_form_decode_failure_forces_ocr_and_preserves_footer(self):
        def damaged_native_pages(*_args, **kwargs):
            kwargs["diagnostic_callback"](0, "Camada nativa incompleta: XForm /Xf1")
            yield "Rodapé técnico do PJe"

        with mock.patch.object(core_pipeline, "extract_text_pages", side_effect=damaged_native_pages), \
             mock.patch.object(core_pipeline, "convert_from_path", return_value=["render.jpg"]), \
             mock.patch.object(core_pipeline, "ocr_image_file", return_value="Conteúdo principal recuperado por OCR"):
            meta, _ = app.process_file(
                self.pdf, self.output_dir, self.input_dir, use_ocr=True, output_format="md", workers=1
            )
        content = Path(meta["output_path"]).read_text(encoding="utf-8")
        self.assertIn("Rodapé técnico do PJe", content)
        self.assertIn("Conteúdo principal recuperado por OCR", content)
        self.assertEqual(meta["page_details"][0]["source"], "nativo+ocr")
        self.assertIn("camada nativa incompleta", meta["page_details"][0]["ocr_reason"])

    def test_direct_image_uses_vision_and_preserves_original(self):
        original = self.image.read_bytes()
        with mock.patch.object(vision_ocr, "is_vision_ocr_available", return_value=True), \
             mock.patch.object(vision_ocr, "ocr_image_vision", return_value="TEXTO DO PRINT"):
            meta, _ = app.process_file(
                self.image, self.output_dir, self.input_dir, use_ocr=True, ocr_backend="auto", output_format="md"
            )
        self.assertEqual(self.image.read_bytes(), original)
        self.assertIn("TEXTO DO PRINT", Path(meta["output_path"]).read_text(encoding="utf-8"))

    def test_tesseract_direct_image_mock(self):
        with mock.patch.object(core_ocr, "pytesseract") as tess:
            tess.image_to_string.return_value = "OCR TESSERACT"
            tess.pytesseract.tesseract_cmd = "tesseract"
            meta, _ = app.process_file(
                self.image,
                self.output_dir,
                self.input_dir,
                use_ocr=True,
                ocr_backend="tesseract",
                tesseract_path="tesseract",
                output_format="txt",
            )
        self.assertIn("OCR TESSERACT", Path(meta["output_path"]).read_text(encoding="utf-8"))

    def test_docx_and_markdown_inputs(self):
        for source, expected in ((self.docx, "Titulo DOCX"), (self.md, "Conteudo integral")):
            meta, _ = app.process_file(source, self.output_dir, self.input_dir, use_ocr=False, output_format="md")
            self.assertIn(expected, Path(meta["output_path"]).read_text(encoding="utf-8"))

    def test_page_audit_records_images_and_ocr_reason(self):
        with mock.patch.object(core_pipeline, "convert_from_path", return_value=["render.jpg"]), \
             mock.patch.object(core_pipeline, "ocr_image_file", return_value="Texto no print"):
            meta, _ = app.process_file(
                self.mixed_pdf,
                self.output_dir,
                self.input_dir,
                use_ocr=True,
                page_audit_enabled=True,
                workers=1,
            )
        audit = Path(meta["page_audit_path"]).read_text(encoding="utf-8")
        self.assertIn('"image_count": 1', audit)
        self.assertIn('"ocr_image_candidate": true', audit)
        self.assertIn('"ocr_reason": "imagem relevante detectada', audit)

    def test_completed_run_cleans_checkpoint(self):
        meta, _ = app.process_file(self.pdf, self.output_dir, self.input_dir, use_ocr=True, output_format="md")
        self.assertEqual(meta["status"], "ok")
        self.assertFalse((self.output_dir / ".hollyocr_checkpoints").exists())

    def test_corrupt_pdf_is_failure(self):
        broken = self.input_dir / "broken.pdf"
        broken.write_bytes(b"%PDF-1.4\ninvalid")
        meta, _ = app.process_file(broken, self.output_dir, self.input_dir, use_ocr=False)
        self.assertEqual(meta["status"], "failed")
        broken.unlink()
        self.assertFalse(broken.exists())

    def test_output_is_atomic_and_unique(self):
        target = self.output_dir / "same.md"
        target.write_text("old", encoding="utf-8")
        self.assertEqual(app.unique_output_path(target).name, "same_2.md")
        out = self.output_dir / "atomic.md"
        app.write_output_document_atomic(
            out, self.pdf, ["texto"], [{"page": 1, "summary": "texto"}], output_format="md"
        )
        self.assertTrue(out.exists())
        self.assertFalse(list(out.parent.glob(f".{out.name}.*.tmp")))

    def test_extraction_api_is_lazy(self):
        pages = extract_text_pages(self.pdf, output_format="txt")
        self.assertFalse(isinstance(pages, list))
        self.assertIn("Documento A", next(iter(pages)))

    def test_automatic_profiles_cover_50_to_20000_pages(self):
        expected = {50: ("pequeno", 50), 1000: ("grande", 32), 5000: ("muito-grande", 24), 20000: ("gigante", 16)}
        for pages, (name, batch) in expected.items():
            profile = choose_processing_profile(pages, 200)
            self.assertEqual((profile.name, profile.render_batch_size), (name, batch))
            self.assertGreaterEqual(profile.ocr_dpi, 300)

    def test_profile_is_reported_in_metadata(self):
        meta, _ = app.process_file(self.pdf, self.output_dir, self.input_dir, use_ocr=False)
        self.assertEqual(meta["processing_profile"], "pequeno")

    def test_process_file_reuses_supplied_executor_instead_of_opening_a_new_pool(self):
        two_page_pdf = self.input_dir / "duas_paginas.pdf"
        doc = fitz.open()
        for text in ("Pagina inicial nativa.", "Segunda pagina nativa."):
            page = doc.new_page()
            page.insert_text((72, 96), text, fontsize=12)
        doc.save(two_page_pdf)
        doc.close()

        with mock.patch.object(core_pipeline, "convert_from_path", return_value=["r1.jpg", "r2.jpg"]), \
             mock.patch.object(core_pipeline, "ocr_image_file", return_value="Texto visto pelo OCR"), \
             mock.patch.object(core_pipeline, "ProcessPoolExecutor") as mocked_pool_cls:
            with ThreadPoolExecutor(max_workers=2) as shared_executor:
                meta, _ = app.process_file(
                    two_page_pdf,
                    self.output_dir,
                    self.input_dir,
                    use_ocr=True,
                    force_ocr=True,
                    output_format="md",
                    workers=2,
                    executor=shared_executor,
                )
        self.assertEqual(meta["status"], "ok")
        mocked_pool_cls.assert_not_called()
        content = Path(meta["output_path"]).read_text(encoding="utf-8")
        self.assertIn("Texto visto pelo OCR", content)

    def test_native_extraction_falls_back_to_pypdf_when_mupdf_unavailable(self):
        """Regression test for the extraction path used by the frozen macOS app,
        where PyMuPDF/pymupdf4llm are disabled for stability (see
        MUPDF_DISABLED_FOR_STABILITY) and every PDF falls back to pypdf."""
        with mock.patch.object(core_extraction, "fitz", None), \
             mock.patch.object(core_extraction, "pymupdf4llm", None):
            pages = list(extract_text_pages(self.pdf, output_format="md"))
        self.assertEqual(len(pages), 1)
        self.assertIn("Documento A", pages[0])


class GuiSharedProcessingTests(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory(prefix="gui_loop_tests_")
        self.base = Path(self.tmp.name)
        self.input_dir = self.base / "input"
        self.output_dir = self.base / "output"
        self.input_dir.mkdir()
        self.output_dir.mkdir()
        self.files = [self.input_dir / "a.md", self.input_dir / "b.md"]
        for idx, path in enumerate(self.files):
            path.write_text(f"# Documento {idx}", encoding="utf-8")

    def tearDown(self):
        self.tmp.cleanup()

    def common(self, stop):
        events = []
        return dict(
            file_list=self.files,
            out_base=self.output_dir,
            input_root=self.input_dir,
            poppler_path=None,
            tesseract_path=None,
            lang="por",
            ocr_threshold=30,
            ocr_dpi=300,
            use_ocr=False,
            force_ocr=False,
            page_audit_enabled=False,
            output_format="md",
            output_mode="individual",
            workers=1,
            should_stop=lambda: stop["value"],
            describe_file_progress=lambda i, total, path: f"{i}/{total} {path.name}",
            make_status_callback=lambda *_: events.append,
            on_file_status=lambda path, status: events.append((path, status)),
            on_progress=lambda value: events.append(("progress", value)),
            on_status=lambda value: events.append(("status", value)),
        ), events

    def test_runs_all_files_and_reaches_100(self):
        kwargs, events = self.common({"value": False})
        metas = gui_shared_processing.run_conversion_loop(**kwargs)
        self.assertEqual(len(metas), 2)
        self.assertIn(("progress", 100), events)

    def test_cancellation_stops_before_work(self):
        kwargs, events = self.common({"value": True})
        metas = gui_shared_processing.run_conversion_loop(**kwargs)
        self.assertEqual(metas, [])
        self.assertTrue(any("Cancelado" in item[1] for item in events if item[0] == "status"))


if __name__ == "__main__":
    unittest.main()
